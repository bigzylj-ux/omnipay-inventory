#!/usr/bin/env python3
"""
Generate comprehensive OmniPay POS Inventory System documentation in DOCX format.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

def add_heading_with_spacing(doc, text, level, space_after=12):
    """Add heading with consistent formatting"""
    heading = doc.add_heading(text, level=level)
    heading.space_after = Pt(space_after)
    return heading

def add_colored_heading(doc, text, level, color_rgb=(16, 120, 80)):
    """Add heading with custom color"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color_rgb)
    return heading

def shade_cell(cell, color):
    """Add background color to table cell"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_table_with_style(doc, rows, cols, headers=None):
    """Create table with styling"""
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light Grid Accent 1'
    
    if headers:
        header_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            shade_cell(cell, 'E8F3EF')
            cell.text = header_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(16, 120, 80)
    
    return table

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ==================== TITLE PAGE ====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run('OmniPay')
title_run.font.size = Pt(48)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(16, 120, 80)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run('POS Inventory System')
subtitle_run.font.size = Pt(28)
subtitle_run.font.color.rgb = RGBColor(71, 85, 119)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_run = info.add_run('Comprehensive User & Management Guide')
info_run.font.size = Pt(14)
info_run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Version and date
version_info = doc.add_paragraph()
version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
version_info.add_run(f'Version: 1.0.0\n')
version_info.add_run(f'Released: {datetime.now().strftime("%B %d, %Y")}\n')
version_info.add_run('Local Mode - Browser-Based Storage')

# Add page break
doc.add_page_break()

# ==================== TABLE OF CONTENTS ====================
add_colored_heading(doc, 'Table of Contents', 1)

toc_items = [
    '1. Executive Summary',
    '2. System Overview & Architecture',
    '3. Technology Stack & Framework',
    '4. Core Components',
    '5. User Guide & Workflows',
    '6. Daily Operations Guide',
    '7. Troubleshooting & Error Resolution',
    '8. Appendix',
]

for item in toc_items:
    p = doc.add_paragraph(item, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)

doc.add_page_break()

# ==================== 1. EXECUTIVE SUMMARY ====================
add_colored_heading(doc, '1. Executive Summary', 1)

doc.add_heading('Project Overview', 2)
summary_text = """
The OmniPay POS Inventory System is a comprehensive, browser-based application designed to manage and track Point-of-Sale (POS) terminal inventory across multiple locations and vendors. The system enables efficient daily reconciliation of inventory against portal records, vendor management, and repair tracking.

Key Capabilities:
• Master inventory management with detailed tracking of terminal status
• Daily automated reconciliation between inventory and portal data
• Vendor management with repair cost tracking
• Real-time search and filtering across all data
• Local persistence ensuring data availability offline
• Export capabilities for reporting and analysis
"""
doc.add_paragraph(summary_text)

doc.add_heading('Key Features', 2)
features = [
    'Dashboard with KPI cards showing inventory status overview',
    'Master inventory search and filtering by serial number, terminal ID, merchant, status, location, and category',
    'Daily portal DB import with automatic deduplication and validation',
    'Automated daily reconciliation with auto-update of terminal assignments',
    'Vendor management including contact details and repair cost tracking',
    'Faulty and under-repair item tracking',
    'Historical reconciliation logs and audit trail',
    'Template download for standardized data import',
    'Excel export functionality for reporting',
]

for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('Target Users', 2)
doc.add_paragraph("""
• Inventory Managers: Oversee daily import and reconciliation processes
• Operations Staff: Handle vendor interactions and repair tracking
• Finance Team: Export reports for analysis and auditing
• New Employees: Can quickly learn the system via intuitive interface and this guide
""")

doc.add_page_break()

# ==================== 2. SYSTEM OVERVIEW ====================
add_colored_heading(doc, '2. System Overview & Architecture', 1)

doc.add_heading('High-Level Architecture', 2)
doc.add_paragraph("""
The OmniPay system follows a modern web architecture:

User Interface Layer (React + TypeScript)
↓
State Management & Routing (React Router)
↓
Data Persistence (IndexedDB - Browser Local Storage)

This architecture ensures:
• No reliance on external servers for data storage
• Full data availability across browser sessions
• Fast performance with instant data retrieval
• Foundation for future server-backed sync
""")

doc.add_heading('Core Workflow', 2)
workflow_steps = [
    ('Master Inventory Setup', 'Upload the initial Data sheet containing all POS terminals'),
    ('Daily Portal Import', 'Download and import the portal DB sheet daily'),
    ('Automatic Reconciliation', 'System matches portal records against inventory'),
    ('Auto-Update', 'Terminal assignments and status updated based on portal data'),
    ('Vendor Tracking', 'Manage vendor repair records and fault costs'),
    ('Reporting', 'Export data for analysis and audit trails'),
]

for step, description in workflow_steps:
    p = doc.add_paragraph()
    p.add_run(step + ': ').bold = True
    p.add_run(description)

doc.add_page_break()

# ==================== 3. TECHNOLOGY STACK ====================
add_colored_heading(doc, '3. Technology Stack & Framework', 1)

tech_table = create_table_with_style(doc, 7, 3, headers=['Layer', 'Technology', 'Purpose'])
tech_data = [
    ['Frontend Framework', 'React 18.2', 'UI components and state management'],
    ['Language', 'TypeScript 5.3', 'Type-safe code development'],
    ['Build Tool', 'Vite 5.0', 'Fast bundling and hot reload'],
    ['Styling', 'Tailwind CSS 3.3', 'Responsive and modern UI design'],
    ['Routing', 'React Router v6', 'Client-side page navigation'],
    ['Data Persistence', 'IndexedDB', 'Local browser-based storage'],
]

for i, row_data in enumerate(tech_data, 1):
    row_cells = tech_table.rows[i].cells
    for j, cell_data in enumerate(row_data):
        row_cells[j].text = cell_data

doc.add_heading('Data Format', 2)
doc.add_paragraph("""
• Import Format: Excel (.xlsx), CSV compatible
• Master Data Sheet: Named 'Data' sheet in workbook
• Portal Data Sheet: Named 'db' sheet in workbook
• Export Format: Excel (.xlsx) for reports and backups
""")

doc.add_page_break()

# ==================== 4. CORE COMPONENTS ====================
add_colored_heading(doc, '4. Core Components', 1)

components = [
    {
        'name': 'Dashboard',
        'desc': 'Real-time overview of inventory status with KPI cards showing total inventory, deployed terminals, terminals yet to deploy, in-stock, faulty items, test terminals, retrieved items, and under repair items. Filterable by category, month, and year.',
    },
    {
        'name': 'Inventory Page',
        'desc': 'Complete searchable inventory with advanced filtering. Search across serial numbers, terminal IDs, merchant names, and phone numbers. Filter by status, location, and category. Export filtered results to Excel.',
    },
    {
        'name': 'Daily Import',
        'desc': 'Upload master inventory (Data sheet) or daily portal DB (db sheet). Automatic deduplication, validation, and meaningful data filtering. Template download for standardized import format.',
    },
    {
        'name': 'Reconciliation Engine',
        'desc': 'Automated matching of portal records against master inventory. Auto-updates terminal assignments, merchant names, phone numbers. Flags exceptions for manual review. Generates detailed audit logs.',
    },
    {
        'name': 'Vendor Management',
        'desc': 'Add and manage vendor contact details (email, phone, address). Upload vendor repair files with fault-based costs. Search and filter faulty/under-repair items. Update terminal status to "Repaired" or "Cannibalised".',
    },
    {
        'name': 'Tracking Page',
        'desc': 'Historical view of terminal assignments and location changes over time. Audit trail for compliance and investigation purposes.',
    },
]

for comp in components:
    heading = doc.add_heading(comp['name'], 3)
    doc.add_paragraph(comp['desc'])

doc.add_page_break()

# ==================== 5. USER GUIDE ====================
add_colored_heading(doc, '5. User Guide & Workflows', 1)

doc.add_heading('5.1 Dashboard Overview', 2)
doc.add_paragraph("""
The Dashboard provides an at-a-glance view of your POS terminal inventory status.

Key Metrics Displayed:
• Total Inventory: Count of all terminals in the system
• Deployed: Terminals actively in use at merchant locations
• Yet To Deploy: Terminals available but not assigned
• In Stock: Terminals held in warehouse/storage
• Faulty: Terminals reported as defective
• Test Terminals: Terminals used for internal testing
• Retrieved: Terminals recovered from field
• Under Repair: Terminals sent to vendors for repair

Use the Dashboard Filters to view data by:
• Category (All, specific categories)
• Month (Current month or historical)
• Year (By fiscal year)
""")

doc.add_heading('5.2 Inventory Management', 2)
doc.add_paragraph("""
Step 1: Access the Inventory Page
Click "Inventory" in the left navigation menu.

Step 2: Search & Filter
• Use the search bar to find terminals by:
  - Device Serial Number
  - Terminal ID
  - Merchant Name
  - Phone Number
• Apply filters for Status, Location, and Category
• Click "Export" to download filtered data as Excel

Step 3: View Details
Hover over any row to see complete terminal information including:
• Device serial number and terminal ID
• Assigned merchant and phone number
• Deployment date and current status
• Fault description and location
• Manager and region assignment
""")

doc.add_heading('5.3 Daily Import Process', 2)
doc.add_paragraph("""
Initial Setup (One-time):

1. Download the Master Inventory Template
   - Click "Download Master Inventory Template" button
   - This provides the correct column structure

2. Prepare Your Data Sheet
   - Use the template columns: SN, Device Serial No., TerminalID, etc.
   - Fill in your terminal data
   - Save as Excel or CSV

3. Upload Master Inventory
   - Go to "Daily Import" page
   - Select "Import Master Inventory (Data Sheet)" tab
   - Click upload area to select file
   - Wait for import confirmation
   - System shows: new records count, updated records count

Daily Workflow:

1. Download Portal DB
   - Get the daily DB sheet from your portal
   - Contains current terminal assignments and status

2. Import Portal DB
   - Go to "Daily Import" page
   - Select "Import Portal DB (Daily Download)" tab
   - Upload the DB file
   - System automatically:
     * Deduplicates records by serial number
     * Filters out empty rows
     * Keeps records with meaningful data
     * Displays import stats

3. Proceed to Reconciliation
   - See Reconciliation section below
""")

doc.add_heading('5.4 Reconciliation Process', 2)
doc.add_paragraph("""
After importing daily portal data, run reconciliation:

1. Go to "Reconciliation" page
2. Review portal records count (shown in green indicator)
3. Click "Run Daily Reconciliation" button
4. System automatically:
   - Matches portal serials against master inventory
   - Auto-assigns terminal IDs where available
   - Updates merchant names from portal
   - Updates phone numbers
   - Flags exceptions (mismatches, missing records)
5. Review results in KPI cards:
   - Records Updated: Auto-matched and updated
   - New Deployments: New terminals found in portal
   - No Change: Terminals with no updates
   - Exceptions Flagged: Issues requiring manual review
6. Export reconciliation logs for audit trail
""")

doc.add_heading('5.5 Vendor Management', 2)
doc.add_paragraph("""
Adding a Vendor:

1. Go to "Vendors" page
2. In "Add Vendor" section, fill in:
   - Vendor Name (required)
   - Email Address (optional)
   - Phone Number (optional)
   - Business Address (optional)
3. Click "Add Vendor" button
4. Vendor appears in the vendors list

Uploading Repair Records:

1. Select vendor from the list
2. In "Vendor Repair Upload" section:
   - Upload Excel file with repair data
   - File must include: Serial Number, Fault columns, Cost
3. System imports and displays repair history
4. View imported repair rows in "Repair History" table

Updating Terminal Status:

1. In "Faulty & Under Repair Inventory" section
2. Find terminal by serial number (search available)
3. Choose new status from dropdown:
   - Repaired: Mark repair as complete
   - Cannibalised: Salvage parts from terminal
4. Status updates immediately in inventory
""")

doc.add_page_break()

# ==================== 6. DAILY OPERATIONS ====================
add_colored_heading(doc, '6. Daily Operations Guide', 1)

doc.add_heading('Standard Daily Checklist', 2)

checklist = [
    'Download portal DB sheet from your portal system',
    'Open OmniPay system (localhost:5174)',
    'Navigate to Daily Import page',
    'Upload the portal DB file',
    'Go to Reconciliation page',
    'Click "Run Daily Reconciliation"',
    'Review reconciliation results (updated, new, no change, exceptions)',
    'If exceptions found, review and manually correct if needed',
    'Export reconciliation logs for records',
    'Optional: Export inventory for reporting',
]

for i, item in enumerate(checklist, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(item)

doc.add_heading('Weekly Tasks', 2)
weekly = [
    'Export complete inventory report for management review',
    'Export repair history from Vendors page',
    'Review any flagged exceptions from reconciliations',
    'Update vendor contact details if changed',
    'Verify all "Under Repair" terminals have been tracked',
]

for task in weekly:
    doc.add_paragraph(task, style='List Bullet')

doc.add_heading('Monthly Tasks', 2)
monthly = [
    'Archive reconciliation logs for audit trail',
    'Review and update vendor performance metrics',
    'Generate monthly inventory variance report',
    'Reconcile inventory counts with physical audit',
    'Update master inventory with any new terminal batches',
]

for task in monthly:
    doc.add_paragraph(task, style='List Bullet')

doc.add_page_break()

# ==================== 7. TROUBLESHOOTING ====================
add_colored_heading(doc, '7. Troubleshooting & Error Resolution', 1)

issues = [
    {
        'problem': 'Import file not recognized',
        'cause': 'Missing expected column headers or wrong sheet name',
        'solution': [
            'For Master Inventory: Ensure sheet is named "Data" or contains "Device Serial No." header',
            'For Portal DB: Ensure sheet is named "db" or contains "Serial Number" header',
            'Use the provided template by clicking "Download Master Inventory Template"',
            'Verify Excel file format (.xlsx or .xls)',
        ]
    },
    {
        'problem': 'Portal records not found when running reconciliation',
        'cause': 'Portal DB file was not imported before reconciliation',
        'solution': [
            'Go back to Daily Import page',
            'Select "Import Portal DB (Daily Download)" tab',
            'Upload your portal DB file',
            'Wait for success message',
            'Then return to Reconciliation page',
        ]
    },
    {
        'problem': 'Data disappears after browser reload',
        'cause': 'Browser storage was cleared or cache issue',
        'solution': [
            'Data is stored in browser IndexedDB - clearing cache may delete it',
            'To prevent data loss, regularly export inventory to Excel backup',
            'Use Chrome DevTools > Application > IndexedDB to verify data presence',
            'Do not clear browser cache if you want to retain data',
        ]
    },
    {
        'problem': 'Reconciliation flags many exceptions',
        'cause': 'Serial number mismatches between portal and inventory',
        'solution': [
            'Review master inventory to ensure accurate serial numbers',
            'Check portal file for data quality issues',
            'Verify serial number format is consistent (case-sensitive)',
            'Run manual verification of mismatched records',
        ]
    },
    {
        'problem': 'Cannot add vendor - form not submitting',
        'cause': 'Browser JavaScript error or form validation',
        'solution': [
            'Open browser DevTools (F12)',
            'Check Console tab for error messages',
            'Ensure Vendor Name field is filled',
            'Try refreshing the page',
            'Clear browser cache and reload',
        ]
    },
    {
        'problem': 'Search not returning results',
        'cause': 'No data imported yet or search term not matching',
        'solution': [
            'Import inventory first via Daily Import > Master Inventory tab',
            'Search is case-insensitive but requires exact field matching',
            'Try searching by different field (serial number vs merchant name)',
            'Verify data was imported by checking Dashboard KPI count',
        ]
    },
    {
        'problem': 'Export to Excel fails or produces empty file',
        'cause': 'No records match current filters or browser permissions issue',
        'solution': [
            'Clear filters to ensure records are visible',
            'Check that records exist in the system',
            'Grant browser permission to download files',
            'Try exporting from Dashboard instead of Inventory page',
        ]
    },
    {
        'problem': 'Slow performance or loading delays',
        'cause': 'Large data volume in browser storage',
        'solution': [
            'Browser performance degrades with very large datasets (>10,000 records)',
            'Consider clearing old reconciliation logs periodically',
            'Use filtering to work with smaller subsets of data',
            'Try accessing the system in a different browser',
        ]
    },
]

for i, issue in enumerate(issues, 1):
    heading = doc.add_heading(f'Issue {i}: {issue["problem"]}', 2)
    
    p = doc.add_paragraph()
    p.add_run('Cause: ').bold = True
    p.add_run(issue['cause'])
    
    doc.add_paragraph('Resolution:', style='List Bullet')
    for solution in issue['solution']:
        doc.add_paragraph(solution, style='List Number')

doc.add_heading('Contact Support', 2)
doc.add_paragraph("""
If you encounter issues not covered in this guide:
1. Check your browser console for error messages (F12 key)
2. Clear browser cache and try again
3. Try the system in a different browser
4. Ensure you have adequate disk space for local storage
5. Document the error with screenshots for support team
""")

doc.add_page_break()

# ==================== 8. APPENDIX ====================
add_colored_heading(doc, '8. Appendix', 1)

doc.add_heading('Master Inventory Template Columns', 2)
doc.add_paragraph('When uploading master inventory, include these columns in order:')

cols = [
    'SN', 'Device Serial No.', 'TerminalID', 'TransactingTid', 'MerchantName',
    'PhoneNo', 'DateMapped', 'SimSerial', 'DateDispatched', 'Custodian',
    'PickupStaff', 'RedispatchMfc', 'Location', 'Status', 'Fault', 'Category', 'Manager', 'Region'
]

for i, col in enumerate(cols, 1):
    p = doc.add_paragraph()
    p.add_run(f'Column {i}: ').bold = True
    p.add_run(col)

doc.add_heading('Terminal Status Types', 2)
doc.add_paragraph('Supported terminal status values:')

statuses = [
    'Deployed - Terminal actively in use',
    'In Stock - Terminal in warehouse',
    'Yet To Deploy - Available for deployment',
    'Under Repair - Sent to vendor for repair',
    'Repaired - Repair completed',
    'Faulty - Terminal has defects',
    'Retrieved - Recovered from field',
    'Cannibalised - Parts salvaged',
    'Lost - Terminal missing',
    'Test Terminal - Used for testing',
]

for status in statuses:
    doc.add_paragraph(status, style='List Bullet')

doc.add_heading('Keyboard Shortcuts', 2)
doc.add_paragraph("""
• F12: Open Browser Developer Tools (for troubleshooting)
• Ctrl+Shift+Delete: Clear browser cache (use carefully - will delete data)
• Ctrl+F: Find within page
• Ctrl+P: Print or save as PDF
""")

doc.add_heading('Data Retention', 2)
doc.add_paragraph("""
• Master Inventory: Persists in IndexedDB until cleared
• Portal Records: Latest import stored, previous imports cleared
• Reconciliation Logs: Last 500 logs kept, older ones auto-removed
• Vendor Data: Persists until manually deleted
• Repair Records: Archived until user clears

Backup Strategy:
• Export inventory weekly to Excel backup files
• Archive reconciliation logs monthly
• Date-stamp all exported files for audit trail
""")

doc.add_heading('System Requirements', 2)
doc.add_paragraph("""
• Modern Web Browser: Chrome, Firefox, Edge, Safari (recent versions)
• Internet Connection: Required for initial access, optional after
• Disk Space: Minimum 50MB for full database
• JavaScript: Must be enabled in browser
• Cookies/Local Storage: Must be enabled in browser
""")

doc.add_heading('FAQ', 2)
faq_items = [
    ('Q: Do I need internet to use the system?',
     'A: After loading the page, the system works offline. Data syncs when connection is restored.'),
    ('Q: Where is my data stored?',
     'A: All data is stored locally in your browser\'s IndexedDB, not on a server.'),
    ('Q: Can multiple people use the same system?',
     'A: Currently, the system is single-user per browser. A future version will support multi-user with server backend.'),
    ('Q: How do I backup my data?',
     'A: Regularly export inventory and reconciliation logs to Excel using the Export buttons.'),
    ('Q: What happens if I clear browser cache?',
     'A: All data will be deleted. Always export backups first.'),
    ('Q: Can I import from other systems?',
     'A: Yes, format your data to match the template columns and import as Excel.'),
]

for q, a in faq_items:
    p = doc.add_paragraph()
    p.add_run(q + '\n').bold = True
    p.add_run(a)

# ==================== FINAL PAGE ====================
doc.add_page_break()

footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.add_run('End of Documentation')
footer_run.font.italic = True
footer_run.font.size = Pt(10)
footer_run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

footer_info = doc.add_paragraph()
footer_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_info.add_run(f'Document Generated: {datetime.now().strftime("%B %d, %Y at %H:%M:%S")}\n')
footer_info.add_run('OmniPay POS Inventory System v1.0.0\n')
footer_info.add_run('Local Mode - Browser-Based Storage')

# Save document
output_path = r'c:\Users\Sunday Aderibigbe\omnipay-inventory\OmniPay_System_Documentation.docx'
doc.save(output_path)
print(f"✓ Documentation generated successfully!")
print(f"✓ File saved to: {output_path}")
print(f"✓ Total pages: Approximately 28-30 pages")
